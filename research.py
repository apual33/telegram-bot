import asyncio
import logging
import re
import smtplib
from datetime import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO
from zoneinfo import ZoneInfo

from anthropic import AsyncAnthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)
_TZ = ZoneInfo("Europe/Berlin")

_SYSTEM = (
    "You are a professional research analyst. The user wants you to research a topic thoroughly.\n\n"
    "You MUST perform at least 5 sequential web searches before writing the report:\n"
    "1. A broad overview of the topic\n"
    "2. Key facts, data, and statistics\n"
    "3. Recent developments (last 1-2 years)\n"
    "4. Expert opinions, criticism, or controversies\n"
    "5. Practical implications or outlook\n"
    "Use findings from each search to sharpen the next query.\n\n"
    "After completing all searches, write a comprehensive structured report. "
    "Format it with these exact markers:\n"
    "  '## ' for main section headers\n"
    "  '### ' for subsection headers\n"
    "  '- ' for bullet points\n\n"
    "Required sections: Executive Summary, Background, Key Findings, Analysis, Conclusion.\n"
    "Minimum 600 words. Cite specific facts found during research."
)

# Detect research intent
_RESEARCH_RE = re.compile(
    r"\b(research|recherchier[e]?|recherche\s+(?:zu\b|über\b|uber\b)?)",
    re.IGNORECASE,
)

def is_research_request(text: str) -> bool:
    return bool(_RESEARCH_RE.search(text))


def extract_topic(text: str) -> str:
    # 1. Explicit question marker: "Frage: ..." or "Question: ..."
    m = re.search(r'\b(?:Frage|Question)\s*:\s*(.+)', text, re.IGNORECASE | re.DOTALL)
    if m:
        return m.group(1).strip().rstrip(".,!?")

    # 2. Verb form: "recherchiere/research [über/zu] X [und schick...]"
    m = re.search(
        r"\b(?:research|recherchier[e]?)\s+(?:über\s+|uber\s+|zu\s+)?(.+?)(?:\s+(?:and|und)\s+send\b|\s+und\s+schick|\s*$)",
        text, re.IGNORECASE,
    )
    if m:
        return m.group(1).strip().rstrip(".,!?")

    # 3. Noun form: "Recherche über/zu/: X"
    m = re.search(
        r"\bRecherche\s*(?:über\s+|uber\s+|zu\s+|:\s*)(.+?)(?:\s+und\s+schick|\s*$)",
        text, re.IGNORECASE | re.DOTALL,
    )
    if m:
        return m.group(1).strip().rstrip(".,!?")

    return text.strip()[:80]


async def run(user_message: str, client: AsyncAnthropic) -> str:
    """Run multi-step web research and return the full report text."""
    from ai import MODEL

    tools = [{"type": "web_search_20250305", "name": "web_search"}]
    messages = [{"role": "user", "content": user_message}]

    for _ in range(40):
        response = await asyncio.wait_for(
            client.messages.create(
                model=MODEL,
                max_tokens=8096,
                system=_SYSTEM,
                tools=tools,
                messages=messages,
                extra_headers={"anthropic-beta": "web-search-2025-03-05"},
            ),
            timeout=120,
        )

        assistant_content = [b.model_dump() for b in response.content]
        messages.append({"role": "assistant", "content": assistant_content})

        if response.stop_reason == "end_turn":
            text_parts = [b.text for b in response.content if b.type == "text"]
            return "\n".join(text_parts) or "(no report generated)"

        # stop_reason == "tool_use": server-side web_search is running.
        # No client execution needed — just continue the loop so the model
        # can process search results and either search again or write the report.
        has_client_tool_use = any(
            b.type == "tool_use" for b in response.content
        )
        if has_client_tool_use:
            # Unexpected client tool call during research — return what we have
            text_parts = [b.text for b in response.content if b.type == "text"]
            return "\n".join(text_parts) or "(research incomplete)"

    return "(research timed out)"


def build_docx(topic: str, report_text: str) -> BytesIO:
    doc = Document()

    title = doc.add_heading(f"Research Report: {topic}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        datetime.now(_TZ).strftime("Generated: %Y-%m-%d %H:%M %Z")
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    for line in report_text.splitlines():
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def send_email(topic: str, docx_buf: BytesIO, filename: str, config) -> None:
    msg = MIMEMultipart()
    msg["From"] = config.gmail_address
    msg["To"] = config.report_email
    msg["Subject"] = f"Research Report: {topic}"

    body = (
        f"Please find attached the research report on: {topic}\n\n"
        f"Generated: {datetime.now(_TZ).strftime('%Y-%m-%d %H:%M %Z')}"
    )
    msg.attach(MIMEText(body, "plain"))

    part = MIMEBase("application", "octet-stream")
    part.set_payload(docx_buf.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{filename}"')
    msg.attach(part)

    try:
        with smtplib.SMTP("smtp.gmail.com", 587, timeout=30) as smtp:
            smtp.starttls()
            smtp.login(config.gmail_address, config.gmail_app_password)
            smtp.sendmail(config.gmail_address, config.report_email, msg.as_string())
    except Exception:
        logging.exception("Email send failed")
